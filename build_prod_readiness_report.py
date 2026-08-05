from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "MacClippy-Prod-Readiness-Report.docx"

# Resolved design tokens: decision_memo (standard_business_brief + Arial).
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}

# Use an installed macOS CJK face as the document face so LibreOffice/Word
# renders Chinese and Latin text consistently in the final artifact. The
# decision-memo hierarchy, spacing, and color tokens remain unchanged.
FONT = "Heiti SC"
# Keep the East Asian face identical so the renderer does not split a single
# mixed Chinese/Latin run across incompatible font metadata.
CJK_FONT = "Heiti SC"
MONO_FONT = "Menlo"
INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
PALE_RED = "FDECEC"
PALE_GOLD = "FFF6D9"
PALE_GREEN = "EAF5EE"
RED = "9B1C1C"
GOLD = "7A5A00"
GREEN = "246B45"
WHITE = "FFFFFF"
BORDER = "CDD5DF"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), CJK_FONT if name == FONT else name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, margins=CELL_MARGINS):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for side, value in margins.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="6"):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    assert sum(widths) == CONTENT_WIDTH_DXA, (widths, sum(widths))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        trpr = row._tr.get_or_add_trPr()
        if trpr.find(qn("w:cantSplit")) is None:
            trpr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[idx]))
            tcw.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def set_keep(paragraph, keep_next=False, keep_lines=True):
    ppr = paragraph._p.get_or_add_pPr()
    if keep_next:
        ppr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        ppr.append(OxmlElement("w:keepLines"))


def paragraph_border_bottom(paragraph, color=BLUE, size="12", space="1"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_field(run, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    rpr.append(rfonts)
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_num_definition(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{abstract_id:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(ppr)
    if kind == "bullet":
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), "Arial")
        rfonts.set(qn("w:hAnsi"), "Arial")
        rfonts.set(qn("w:eastAsia"), CJK_FONT)
        rpr.append(rfonts)
        lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)
    ppr.append(numpr)


def style_run(run, style="body"):
    tokens = {
        "body": (FONT, 11, "222222", False, False),
        "small": (FONT, 9, MUTED, False, False),
        "table": (FONT, 9.2, "222222", False, False),
        "table_header": (FONT, 9.2, INK, True, False),
        "label": (FONT, 10.5, INK, True, False),
        "code": (MONO_FONT, 8.6, DARK_BLUE, False, False),
    }
    set_run_font(run, *tokens[style])


def add_text_paragraph(doc, text, style="Normal", before=0, after=6, line=1.10, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    if style == "Caption":
        style_run(run, "small")
    elif style == "Body Small":
        style_run(run, "small")
    elif style == "Code Paragraph":
        style_run(run, "code")
    else:
        style_run(run, "body")
    return p


def add_rich_paragraph(doc, parts, style="Normal", before=0, after=6, line=1.10):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    for text, kind in parts:
        run = p.add_run(text)
        style_run(run, kind)
    return p


def add_bullet(doc, text, num_id, after=8):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.167
    apply_num(p, num_id)
    run = p.add_run(text)
    style_run(run, "body")
    return p


def add_numbered(doc, text, num_id, after=8):
    return add_bullet(doc, text, num_id, after)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_keep(p, keep_next=True)
    return p


def add_table(doc, headers, widths=None, rows=None, header_fill=LIGHT_GRAY, font_size=9.2, first_col_bold=False):
    # Support both the readable ``headers, rows, widths`` call form used for
    # most tables and the ``headers, widths, rows=...`` form used for dense
    # comparison tables.
    if isinstance(widths, list) and widths and isinstance(widths[0], (tuple, list)):
        widths, rows = rows, widths
    if widths is None or rows is None:
        raise ValueError("add_table requires rows and widths")
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    header = table.rows[0]
    repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(text)
        style_run(run, "table_header")
    for row_values in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_values):
            cell = cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(text))
            set_run_font(run, FONT, font_size, "222222", bold=(first_col_bold and idx == 0))
            if first_col_bold and idx == 0:
                run.font.color.rgb = rgb(INK)
    # Geometry must also cover rows created after the initial table.
    set_table_geometry(table, widths)
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return table


def add_callout(doc, label, text, fill=CALLOUT, label_color=INK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=BLUE, size="10")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(label + "  ")
    set_run_font(run, FONT, 10.5, label_color, True, False)
    run = p.add_run(text)
    set_run_font(run, FONT, 10.5, "222222", False, False)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_source_note(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, "small")
    return p


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb("222222")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = FONT
    title._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    title.font.size = Pt(23)
    title.font.bold = True
    title.font.color.rgb = rgb(INK)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.line_spacing = 1.0

    subtitle = styles["Subtitle"]
    subtitle.font.name = FONT
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.line_spacing = 1.05

    heading_tokens = {
        "Heading 1": (16, BLUE, 12, 6),
        "Heading 2": (13, BLUE, 10, 5),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    for name, size, color in (("Caption", 9, MUTED), ("Body Small", 9.2, MUTED), ("Code Paragraph", 8.6, DARK_BLUE)):
        style = styles[name] if name in styles else styles.add_style(name, 1)
        style.font.name = MONO_FONT if name == "Code Paragraph" else FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT if name != "Code Paragraph" else MONO_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.05


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("MacClippy  /  Production Readiness Review")
    set_run_font(r, FONT, 8.5, MUTED, False, False)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    r = fp.add_run("Internal decision brief  ·  Page ")
    set_run_font(r, FONT, 8.5, MUTED, False, False)
    r = fp.add_run()
    set_run_font(r, FONT, 8.5, MUTED, False, False)
    add_field(r, "PAGE")


def add_masthead(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(6)
    run = kicker.add_run("PRODUCTION READINESS REVIEW")
    set_run_font(run, FONT, 9.5, BLUE, True, False)
    title = doc.add_paragraph(style="Title")
    title.add_run("MacClippy 生产就绪性审计与产品化路线图")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("面向稳定运行、隐私信任与正式发布的工程决策备忘录")

    metadata = [
        ("对象", "mac-clippy 独立 macOS 剪贴板管理器"),
        ("日期", "2026 年 8 月 3 日"),
        ("当前判断", "Advanced beta；尚未达到 public-prod-ready"),
        ("核心决策", "继续采用 Swift + AppKit/SwiftUI；暂不整体迁移 Tauri"),
        ("发布前重点", "签名/公证、数据恢复、并发边界、输入上限、隐私默认值、可观测性"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(label + "：")
        set_run_font(r, FONT, 10.5, INK, True, False)
        r = p.add_run(value)
        set_run_font(r, FONT, 10.5, "222222", False, False)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(12)
    paragraph_border_bottom(rule, color=BLUE, size="14", space="1")

    add_callout(
        doc,
        "一句话结论",
        "MacClippy 已经具备先进 beta 的功能完整度和不错的原生 macOS 基础，但还不能以“稳定、可恢复、可持续分发”的 public product 标准发布。下一阶段应优先收敛发布与可靠性风险，而不是重写 UI 技术栈。",
        fill=BLUE_GRAY,
    )


def add_section_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document()
    setup_styles(doc)
    setup_page(doc)
    bullet_id = add_num_definition(doc, "bullet")
    decimal_id = add_num_definition(doc, "decimal")
    add_masthead(doc)

    add_heading(doc, "1. 结论先行", 1)
    add_text_paragraph(doc, "本报告的建议是：保留 Swift + AppKit/SwiftUI 作为 MacClippy 的 macOS 主架构，把接下来的工程投入集中在“可恢复的数据层、串行化的系统边界、可证明的发布链路、隐私安全默认值，以及用户可理解的失败状态”上。Tauri 可以在未来有明确跨平台需求时作为新的 UI/产品边界候选，但不应成为当前稳定性工作的替代方案。")
    add_bullet(doc, "产品状态：功能面已经超过最小可用工具，属于 advanced beta；距离 public-prod-ready 主要差在发布工程和极端场景可靠性，而不是核心功能数量。", bullet_id)
    add_bullet(doc, "架构状态：原生 Swift 对剪贴板、Accessibility、CGEvent、菜单栏、NSPanel、多屏和 TCC 权限拥有最短路径；整体迁移 Tauri 会新增 WebView/Rust/前端桥接边界，不能消除 Apple 签名、公证、权限或 crash 风险。", bullet_id)
    add_bullet(doc, "验证状态：MacClippyKit 179 项测试通过；App XCTest 114 项通过、0 失败、0 跳过。这是很好的回归信号，但还不是对真实 TCC、Keychain、CGEvent、正式签名、多屏和恶意 pasteboard 输入的完整证明。", bullet_id)
    add_bullet(doc, "发布状态：当前 Debug artifact 是 ad-hoc/linker-signed，TeamIdentifier 未设置；尚无被验证的 Developer ID、Hardened Runtime、notarization、stapling 与 CI release pipeline。", bullet_id)
    add_bullet(doc, "最高优先级：建立可签名发布链路、数据库 quick check 与 backup/restore、输入 payload 限制、敏感 pasteboard 默认过滤、structured crash-safe logging，以及真实权限/系统边界测试。", bullet_id)

    add_heading(doc, "2. 现状与证据边界", 1)
    add_text_paragraph(doc, "审计范围聚焦 mac-clippy/ 目录。报告中的代码判断来自仓库源码、构建产物、测试结果、DeepWiki 竞品资料和公开产品页面；外部资料用于提炼可迁移的工程经验，不代表对竞品内部实现的绝对断言。")
    add_heading(doc, "2.1 已经具备的产品基础", 2)
    add_table(
        doc,
        ["能力面", "当前观察", "对产品化的意义"],
        [
            ("捕获与存储", "多 representation 捕获、加密 blob、Keychain 设备密钥、历史与 pinboard/snippet 数据。", "产品核心已成形；下一步应补可恢复性和边界限制。"),
            ("检索", "FTS 搜索、结构化搜索、重复处理、历史浏览。", "效率优势明显；需要补 FTS 修复和降级状态。"),
            ("粘贴", "CGEvent/Accessibility 注入、队列粘贴、失败日志与恢复路径基础。", "体验接近成熟工具；需覆盖键盘布局和失败回滚。"),
            ("交互", "菜单栏常驻、⌘⇧V picker、键盘导航、preview、多屏布局与 dock 动画。", "原生 macOS 体验是重要竞争力；需继续处理状态可见性和权限引导。"),
            ("扩展能力", "OCR、snippet expansion、app/regex 排除、pinboards、转换/预览等。", "差异化空间大；必须用隐私和性能边界保护复杂度。"),
        ],
        [1800, 4300, 3260],
        first_col_bold=True,
    )
    add_source_note(doc, "依据：mac-clippy/README.md、MacClippyRuntime.swift、MacClippyKit 的 Core/Platform 实现及现有测试。")

    add_heading(doc, "2.2 可复现的验证结果", 2)
    add_table(
        doc,
        ["检查项", "结果", "解释与限制"],
        [
            ("MacClippyKit package tests", "179 passed / 0 failed", "覆盖大量纯逻辑、存储和策略回归；不等价于真实系统权限测试。"),
            ("App XCTest（arm64）", "114 passed / 0 failed / 0 skipped", "在本地 macOS 26.5.2 上完成；仍需 Release/TCC/Keychain/CGEvent 真实环境矩阵。"),
            ("Debug 构建产物", "可构建、可运行", "当前为 ad-hoc/linker-signed；不能证明 Developer ID 或公证发布可用。"),
            ("正式发布链路", "未完成证据", "缺 Developer ID、Hardened Runtime、notarization、stapling、CI release gate 的端到端验证。"),
            ("性能/压力", "未完成证据", "尚未用 Instruments 或压力夹具证明超大 pasteboard、多屏、长时间运行和 blob/FTS 维护边界。"),
        ],
        [2500, 2300, 4560],
        first_col_bold=True,
    )
    add_source_note(doc, "证据日期：2026-07-31 审计闭环结果；本报告日期为 2026-08-03。")
    add_callout(doc, "证据边界", "“测试通过”证明回归面稳定，不证明系统边界已经稳定。发布前必须把签名身份、TCC 权限、Keychain、Accessibility、pasteboard provider、CGEvent 注入和异常退出纳入独立的 integration/release test matrix。", fill=PALE_GOLD, label_color=GOLD)

    add_heading(doc, "3. Swift vs. Tauri：决策与边界", 1)
    add_heading(doc, "3.1 决策", 2)
    add_callout(doc, "推荐", "继续用 Swift + AppKit/SwiftUI 做 macOS 产品，不做整体 Tauri rewrite。优先把现有原生路径打磨到可签名、可恢复、可诊断；只有在 Windows/Linux 成为明确商业要求时，再设计平台无关 domain/core 与跨平台 UI 的边界。", fill=PALE_GREEN, label_color=GREEN)
    add_text_paragraph(doc, "Tauri 的价值是跨平台 UI 与 Rust 后端组合，而不是“绕过 macOS”。在 macOS 上，Tauri 仍然需要 Apple 的签名、公证、权限说明和系统级集成；同时会引入 WebView 渲染、Rust/前端消息桥、异步生命周期和额外打包矩阵。对一个高度依赖 clipboard、Accessibility、CGEvent、NSPanel、菜单栏和原生权限的工具，这些新增边界会把当前的稳定性预算分散掉。")

    add_heading(doc, "3.2 关键维度对比", 2)
    add_table(
        doc,
        ["维度", "继续 Swift", "整体迁移 Tauri", "判断"],
        [1700, 2700, 3000, 1960],
        rows=[
            ("macOS 集成", "AppKit/SwiftUI、NSPanel、菜单栏、CGEvent、TCC 路径直接。", "需要 Rust/JS 与原生插件桥接；WebView 适合 UI，不天然适合系统边界。", "Swift 明显更优"),
            ("稳定性风险", "已有代码和测试资产；风险集中在可识别的生命周期和边界。", "重写会同时引入新运行时、桥接和迁移风险。", "Swift 风险更低"),
            ("性能/资源", "原生窗口、列表和事件路径更可控。", "前端 WebView + Rust 通信有额外内存和序列化成本。", "Swift 更匹配"),
            ("跨平台", "需要另做 UI/平台实现。", "一次共享较多 UI/业务代码，Windows/Linux 更有优势。", "Tauri 仅在跨平台承诺成立时占优"),
            ("发布与签名", "需要 Developer ID/公证。", "同样需要 Apple 签名/公证，还多一层打包与 WebView 验证。", "都不能绕过发布工程"),
            ("团队/资产", "现有源码、AppKit 交互和测试全部可复用。", "需要迁移 UI、数据桥和系统集成测试。", "继续 Swift"),
        ],
        header_fill=BLUE_GRAY,
        first_col_bold=True,
    )
    add_source_note(doc, "依据：Tauri 官方架构与 macOS 签名文档；MacClippy 当前 AppKit/SwiftUI、CGEvent、NSPanel 与 TCC 依赖。")

    add_heading(doc, "3.3 如果未来要跨平台，正确的演进方式", 2)
    add_numbered(doc, "先稳定 macOS：把 capture、dedupe、retention、search、paste policy、privacy policy 做成可测试的 platform-neutral domain API。", decimal_id)
    add_numbered(doc, "把平台适配收敛为协议：PasteboardReader、PasteInjector、GlobalHotkey, PermissionProvider、SecureKeyStore、WindowSurface。", decimal_id)
    add_numbered(doc, "定义持久化迁移格式和兼容性版本；不要让未来 Tauri UI 直接读写 SQLite/blob 细节。", decimal_id)
    add_numbered(doc, "只有当第二个平台有真实用户、路线图和维护人力时，才做 Tauri UI spike，并以并行实验替代生产重写。", decimal_id)
    add_numbered(doc, "以体验、内存、冷启动、权限、签名、崩溃率和维护成本做 go/no-go，而不是以“代码可共享比例”单一指标做决定。", decimal_id)

    add_heading(doc, "4. 竞品与可迁移经验", 1)
    add_text_paragraph(doc, "竞品比较的目标不是复制功能，而是识别用户已经形成的期待：快速打开和检索、可信的隐私边界、稳定的粘贴、可组织的历史、成熟的权限/发布体验，以及在极端状态下不丢数据。")
    add_table(
        doc,
        ["产品", "定位/强项", "值得借鉴", "对 MacClippy 的启示"],
        [1500, 2600, 3000, 2260],
        rows=[
            ("Paste", "高完成度商业化跨设备剪贴板。", "跨设备体验、pinboard/组织、多项按顺序粘贴、链接/图片/文件预览、隐私表达。", "把速度、搜索、组织和信任做成一条完整体验；不要只堆功能。"),
            ("Deck", "原生 Swift，SQLite + blob，可选加密/同步，带 CLI/MCP。", "完整性检查、备份恢复、迁移修复、bounded input、hotkey teardown、paste 失败恢复。", "可靠性基础设施是高级用户愿意长期依赖的核心。"),
            ("Maccy", "轻量、开源、原生 macOS、键盘优先。", "changeCount 轮询、concealed/transient/auto-generated 过滤、app/regex 排除、键盘布局处理。", "安全默认值和 picker 的确定性比更多设置更重要。"),
            ("Clipy", "成熟的原生开源剪贴板历史。", "持续集成、发布自动化、迁移、Sparkle、稳定签名对 Accessibility 权限的影响。", "正式签名不是发布最后一步，而是权限体验的一部分。"),
            ("PasteBar", "开源、本地、跨平台/丰富工作流。", "collections/tabs/boards、lock/passcode、模板与工作流。", "可把组织和 workflow 作为差异化，但需控制复杂度与隐私风险。"),
            ("CopyQ", "成熟跨平台、插件和自动化生态。", "长期维护、扩展和跨平台覆盖。", "跨平台的维护面会快速扩大；不要在没有用户拉力时提前承担。"),
        ],
        header_fill=BLUE_GRAY,
        first_col_bold=True,
        font_size=8.9,
    )
    add_source_note(doc, "研究入口：Paste 官方产品页；Deck/Maccy/Clipy/PasteBar DeepWiki；CopyQ GitHub；Tauri 官方文档。")
    add_callout(doc, "竞品结论", "MacClippy 的潜在优势不是“功能最多”，而是“原生 macOS 交互 + 本地隐私 + 深度检索/转换 + 可恢复的数据层”。要赢得长期信任，稳定性、权限身份和隐私默认值必须先达到产品底线。", fill=BLUE_GRAY)

    add_heading(doc, "5. 风险分级与处理顺序", 1)
    add_text_paragraph(doc, "分级标准：P0 是会阻止 public release、造成不可接受的数据/信任损失，或让安装后的核心能力不稳定；P1 是不应长期带入稳定版本的高概率/高影响问题；P2 是体验、效率或后续规模化问题。")
    add_heading(doc, "5.1 P0：发布前必须关闭", 2)
    add_table(
        doc,
        ["风险", "为什么是 P0", "完成标准"],
        [2300, 3560, 3500],
        rows=[
            ("Developer ID / Hardened Runtime / notarization / CI", "当前 Debug 为 ad-hoc；用户无法获得稳定的权限身份和可信安装路径。", "Release 构建由 CI 产出；Developer ID 签名、Hardened Runtime、公证、staple、Gatekeeper 验证全部通过。"),
            ("数据库 quick check + backup/restore", "异常退出、磁盘问题或迁移异常可能让历史/FTS 长期退化而用户不知情。", "启动 quick_check；定期/手动安全备份；演练 restore、迁移失败和回滚；UI 显示修复结果。"),
            ("捕获 payload 硬上限", "当前会先 materialize raw representation，再判断 32 KiB spill；超大/恶意 UTI 可制造内存压力。", "单 representation、单事件、总队列均有上限；超限可预测地丢弃/转 blob；有压力测试和日志指标。"),
            ("敏感 pasteboard 默认过滤", "concealed/transient/auto-generated 和未知 UTI 默认过宽，可能捕获密码、支付信息和一次性秘密。", "默认安全；明确的 Capture All 高级开关；设置页解释规则；敏感数据测试覆盖。"),
            ("crash-safe structured logging + diagnostics", "没有可诊断的失败上下文，用户只能重启或丢失信任。", "日志不含原文/秘密；事件有 correlation ID、阶段、耗时、错误码；可导出脱敏诊断包。"),
        ],
        header_fill=PALE_RED,
        first_col_bold=True,
        font_size=8.9,
    )
    add_heading(doc, "5.2 P1：稳定版本前完成", 2)
    add_table(
        doc,
        ["领域", "问题/改进", "验收信号"],
        [2100, 4240, 3020],
        rows=[
            ("并发边界", "收紧 MacClippyRuntime 的 @unchecked Sendable；PasteboardObserver 的 timer/handler/generation 由单一 lifecycle executor 串行管理。", "TSAN/压力运行无竞态；start/stop/restart 并发测试稳定。"),
            ("lazy provider", "重新读取 unavailable UTI 前验证 changeCount/generation 仍是原始事件；过期 provider 只能被丢弃。", "换 clipboard 后旧 provider 不会污染新记录；回归测试覆盖。"),
            ("paste 失败恢复", "注入失败时明确恢复原 clipboard，区分 injected/manualPasteRequired，并给出可操作错误。", "权限拒绝、目标 app 不响应、队列中断均不会静默破坏用户 clipboard。"),
            ("hotkey 更新", "注册新快捷键失败时回滚到旧配置；卸载/禁用时确定性 teardown。", "冲突、权限、重复注册均有可见状态，旧快捷键不丢。"),
            ("Launch at Login", "移除静默 try?；读取 SMAppService.status 并把失败原因映射到 UI。", "设置值与系统真实状态一致，失败可重试/可诊断。"),
            ("retention", "增加 deletion journal 或可恢复队列，避免 DB/blob 删除半完成时无从追踪。", "强杀/断电模拟后可继续清理，数据库和 blob 一致。"),
            ("FTS repair", "记录成功持久化但 FTS 更新失败的记录，启动/设置页可重建索引。", "搜索完整性可检测、可修复，且不需要删除主记录。"),
        ],
        header_fill=PALE_GOLD,
        first_col_bold=True,
        font_size=8.9,
    )
    add_heading(doc, "5.3 P2：产品化增强", 2)
    add_bullet(doc, "增加跨设备/备份导出、冲突处理和版本化迁移的产品设计；Sync Engine 仍可保持延期，先把本地数据层做到可信。", bullet_id)
    add_bullet(doc, "增加可解释的存储占用视图：历史、blob、缩略图、FTS、备份分别占多少空间，用户可以预览清理影响。", bullet_id)
    add_bullet(doc, "增加自动化 UI smoke：⌘⇧V 打开、首项 focus、Space preview、左右导航、回车粘贴、Esc 关闭、多屏切换。", bullet_id)
    add_bullet(doc, "为收藏夹、pinboard、snippet、结构化搜索增加导入/导出和可迁移 schema，降低用户被单一版本锁定的焦虑。", bullet_id)
    add_bullet(doc, "把性能指标纳入 release dashboard：启动时间、picker 首次可交互、capture p95、paste p95、FTS 查询 p95、内存峰值。", bullet_id)

    add_section_break(doc)
    add_heading(doc, "6. 稳定运行方案：从代码边界到可恢复系统", 1)
    add_heading(doc, "6.1 数据完整性与恢复", 2)
    add_text_paragraph(doc, "剪贴板管理器的核心不是“把内容存进去”，而是让用户相信历史不会因为一次强杀、磁盘满、迁移中断或索引损坏而无声消失。建议把数据库、blob、FTS 和迁移视为一个可观测的数据生命周期。")
    add_bullet(doc, "启动顺序：打开 DB → 执行 quick_check → 检查 schema/migration version → 扫描孤儿 blob/缺失 blob → 校验 FTS 健康 → 再启用 capture observer。", bullet_id)
    add_bullet(doc, "写入顺序：先生成可重试的 record ID 和事务边界；主记录、representation、blob 引用和 FTS 更新要有明确状态。FTS 失败不能被当作只记日志的旁路错误。", bullet_id)
    add_bullet(doc, "恢复工具：提供安全 backup、restore 到临时目录、导入前校验、迁移失败回滚和“仅重建 FTS”操作；所有修复结果都显示数量和错误。", bullet_id)
    add_bullet(doc, "空间治理：分别限制单条记录、单 representation、当天捕获量和总磁盘占用；保留策略、blob 清理和缩略图清理必须幂等。", bullet_id)
    add_bullet(doc, "数据不变式：主记录存在时 representation 引用不能悬空；删除完成后不能残留可恢复原文的 blob；pinboard/snippet 引用必须能处理目标记录已删除。", bullet_id)

    add_heading(doc, "6.2 pasteboard 与异步生命周期", 2)
    add_bullet(doc, "为每次系统 pasteboard change 建立 generation：capture、lazy provider reread、OCR/thumbnail、persist 全部携带 generation，提交前验证仍属于同一代。", bullet_id)
    add_bullet(doc, "PasteboardObserver 的 start/stop/updateHandler/cancelTimer 不要从多个线程直接改共享状态；由 actor、serial executor 或严格的 main-actor 生命周期统一管理。", bullet_id)
    add_bullet(doc, "对 advertised UTI 使用预算式读取：先读 size/metadata 或流式写入 blob，达到上限立即终止；不要在内存中 materialize 未知大小的数据。", bullet_id)
    add_bullet(doc, "区分复制来源：用户复制、应用内部回放、MacClippy 自己的 paste injection、Universal Clipboard/自动生成内容要可追踪并默认采取不同过滤策略。", bullet_id)
    add_bullet(doc, "观察指标：每次 changeCount、忽略原因、代表性读取耗时、超限次数、provider 重读次数、过期 generation 丢弃次数。", bullet_id)

    add_heading(doc, "6.3 粘贴、权限和恢复体验", 2)
    add_bullet(doc, "把 Accessibility 缺失分成“捕获正常但无法自动粘贴”“目标 app 拒绝”“注入超时”“需要手动粘贴”四类，而不是统一显示失败。", bullet_id)
    add_bullet(doc, "paste 前保存用户当前 clipboard 的完整状态；paste 后仅在确认 injection 成功或用户选择手动粘贴时释放临时状态；失败必须恢复。", bullet_id)
    add_bullet(doc, "覆盖非 QWERTY 键盘布局、密码框/安全输入、Electron/虚拟机/远程桌面、Terminal、浏览器地址栏和无焦点窗口。", bullet_id)
    add_bullet(doc, "权限 onboarding 不要只给系统设置链接；解释为什么需要 Accessibility、Input Monitoring、Launch at Login，以及权限被撤回后的恢复方式。", bullet_id)
    add_bullet(doc, "hotkey 注册更新是事务：新 shortcut 成功后再切换；冲突或权限失败时保留旧 shortcut，并把冲突 app/状态显示出来。", bullet_id)

    add_heading(doc, "6.4 可观测性与 crash hardening", 2)
    add_bullet(doc, "使用统一事件模型：subsystem、operation、generation、record ID（不可逆/截断）、duration、result、error code、retry count。", bullet_id)
    add_bullet(doc, "日志禁止写入 clipboard 原文、图片、密码、URL query、Keychain 值；诊断包默认脱敏，用户主动导出前显示包含哪些元数据。", bullet_id)
    add_bullet(doc, "崩溃前写入小型 ring buffer，不在 crash handler 中做复杂 IO；下次启动读取上次异常阶段，提示“上次退出时正在做什么”。", bullet_id)
    add_bullet(doc, "每个异步任务都有 cancellation、deadline、retry policy；不要让 provider、OCR、thumbnail 或 FTS 任务无限等待并堆积。", bullet_id)
    add_bullet(doc, "长时间 soak test：启动后运行 24 小时，持续复制文本/图片/文件、开关 picker、切屏、锁屏/解锁、重启 Finder/目标 app，并记录内存和任务数量。", bullet_id)

    add_heading(doc, "7. 隐私、安全与权限信任", 1)
    add_callout(doc, "产品原则", "剪贴板天然包含密码、token、支付信息和一次性代码。默认策略应该优先保护“用户没有意识到自己正在被记录”的场景；愿意捕获一切的高级用户可以显式打开 Capture All，而不是让普通用户先承担风险。", fill=PALE_RED, label_color=RED)
    add_table(
        doc,
        ["区域", "建议默认", "需要用户看见的说明"],
        [1900, 3560, 3900],
        rows=[
            ("敏感类型", "过滤 concealed、transient、auto-generated、password-like 和已知 token/OTP 来源。", "“哪些内容不会进入历史”；提供 Capture All 的风险提示。"),
            ("应用排除", "支持 bundle ID、应用名、regex；预设密码管理器、银行、终端等高风险场景。", "排除规则是否只影响捕获还是也影响已有历史。"),
            ("本地存储", "原文/图片/blob 加密；Keychain key 失败时不静默降级为明文。", "密钥在哪里、卸载/迁移/备份如何处理。"),
            ("日志诊断", "不记录原文；错误只保留分类、长度、UTI 哈希/计数等必要信息。", "导出诊断包前显示范围，允许取消。"),
            ("权限", "首次使用前解释 Accessibility/Input Monitoring；权限撤回后不反复弹窗。", "权限用途、如何开启、为何仍能查看历史但不能自动粘贴。"),
            ("保留策略", "默认有限保留期/容量，并在设置中展示当前占用。", "删除是永久删除、是否影响 pinboard/snippet、备份是否包含历史。"),
        ],
        header_fill=PALE_RED,
        first_col_bold=True,
        font_size=9,
    )
    add_heading(doc, "7.1 建议的隐私设置分层", 2)
    add_numbered(doc, "安全默认：过滤敏感/瞬态内容，捕获常见文本、图片、文件与链接。", decimal_id)
    add_numbered(doc, "自定义排除：用户按 app、bundle ID、regex 增加规则，实时显示命中次数。", decimal_id)
    add_numbered(doc, "Capture All：明确警告可能记录密码和一次性秘密；要求用户确认，并提供快捷关闭。", decimal_id)
    add_numbered(doc, "Emergency Pause：菜单栏和 picker 都能一键暂停；暂停状态必须持久化并显眼显示。", decimal_id)

    add_heading(doc, "8. UX Polish：从“功能齐全”到“值得长期依赖”", 1)
    add_heading(doc, "8.1 首次运行与权限", 2)
    add_bullet(doc, "首次运行用三句话解释价值：⌘⇧V 打开历史、搜索和 pinboard 快速找回、内容默认只在本机保存。然后再解释权限，不要先展示一串系统术语。", bullet_id)
    add_bullet(doc, "权限状态用明确的动词：允许自动粘贴、去系统设置、重新检查；不要只显示红/绿 pill。", bullet_id)
    add_bullet(doc, "把“捕获已暂停”“搜索索引需要修复”“数据库处于只读恢复模式”“只能手动粘贴”做成统一状态模型，菜单栏、picker、设置页一致。", bullet_id)

    add_heading(doc, "8.2 Picker 与搜索", 2)
    add_bullet(doc, "⌘⇧V 打开后第一项必须有确定的 focus/selection；Space 无 preview target 时仍应被 picker 消费，避免 AppKit beep。", bullet_id)
    add_bullet(doc, "搜索结果把来源、时间、类型、大小和排除/锁定状态变成稳定的辅助信息；匹配片段要突出，但不要把原文塞满卡片。", bullet_id)
    add_bullet(doc, "支持渐进式结果：先显示可用 metadata，再加载图片/文件缩略图；加载失败显示可操作的“重试/打开原文件”。", bullet_id)
    add_bullet(doc, "为“没有结果、搜索索引损坏、历史被暂停、权限未开”分别设计 empty/error/degraded state，避免都显示成空白。", bullet_id)
    add_bullet(doc, "键盘路径应形成闭环：打开 → 上下/左右 → preview → Enter 粘贴 → Esc 关闭；每一步都有焦点反馈和无障碍 label。", bullet_id)

    add_heading(doc, "8.3 多屏、动效与可访问性", 2)
    add_bullet(doc, "保留当前的 display fallback 顺序（cursor screen → 当前 panel screen → main screen），并在屏幕变化时无动画更新最终 frame，避免 dock 飞到错误屏幕。", bullet_id)
    add_bullet(doc, "动效只服务于位置/状态反馈；Reduce Motion 下关闭位移动画，保留即时 opacity/state change。", bullet_id)
    add_bullet(doc, "支持 VoiceOver 的 card role、当前 selection、preview 状态、粘贴结果；键盘焦点与视觉焦点必须一致。", bullet_id)
    add_bullet(doc, "字体放大、深色模式、高对比度和窗口高度受限时，优先压缩次要 metadata，不缩小主要内容到不可读。", bullet_id)

    add_heading(doc, "8.4 错误文案与成功反馈", 2)
    add_table(
        doc,
        ["场景", "不建议", "建议文案/动作"],
        [2200, 2600, 4560],
        rows=[
            ("Accessibility 未开启", "Paste failed", "“历史已复制，但 MacClippy 没有自动粘贴权限。去系统设置开启 Accessibility；也可以按 ⌘V 手动粘贴。”"),
            ("搜索索引失败", "No results", "“历史仍在，但搜索索引需要修复。立即重建索引 / 稍后处理。”"),
            ("内容超限", "Unexpected error", "“这次复制内容过大，未保存到历史。调整大小限制 / 仅手动粘贴。”"),
            ("启动登录失败", "Setting saved", "“系统未接受登录项注册，设置未生效。重新尝试 / 打开系统设置。”"),
            ("数据库恢复", "Loading", "“正在检查历史完整性。捕获已暂停，完成后会自动恢复。”"),
        ],
        header_fill=BLUE_GRAY,
        first_col_bold=True,
        font_size=9,
    )

    add_heading(doc, "9. 4–6 周产品化路线图", 1)
    add_text_paragraph(doc, "假设已有一名主要工程负责人和一名可协作的 UI/QA 负责人；每周都要有可运行 artifact 和可回归的验收结果。若签名证书/Apple 账号尚未就绪，应把第 1 周的发布链路前置为最高依赖。")
    add_table(
        doc,
        ["周期", "主目标", "交付物", "退出条件"],
        [1700, 2300, 3300, 2060],
        rows=[
            ("第 1 周", "发布链路与风险冻结", "Release configuration；Developer ID/Hardened Runtime 签名脚本；CI build/test；崩溃与结构化日志 schema；P0 风险登记。", "CI 能构建 signed artifact；本地 Gatekeeper 验证路径可复现。"),
            ("第 2 周", "数据完整性与恢复", "quick_check；backup/restore；迁移回滚；FTS repair；blob orphan scan；retention deletion journal。", "强杀/断电/磁盘满演练后，历史和搜索可恢复，结果可解释。"),
            ("第 3 周", "pasteboard/并发/输入边界", "generation 校验；observer lifecycle executor；representation/event/queue limits；lazy provider 过期丢弃。", "TSAN/压力测试无已知竞态；超大/未知 UTI 不造成内存失控。"),
            ("第 4 周", "粘贴与权限闭环", "Accessibility/Input Monitoring onboarding；paste 失败恢复；hotkey rollback；Launch at Login 状态同步；键盘布局测试。", "权限拒绝、目标 app 异常、shortcut 冲突都能恢复且不丢用户 clipboard。"),
            ("第 5 周", "隐私默认与 UX polish", "敏感/瞬态过滤；Capture All；Pause；degraded states；picker/search/多屏/Reduce Motion polish。", "隐私设置可理解；⌘⇧V 键盘路径和错误状态可用；无明显 beep/焦点丢失。"),
            ("第 6 周", "Release candidate 与 soak", "24–72 小时 soak；多机器/多屏/锁屏矩阵；公证/staple；用户文档、隐私说明、回滚包。", "Release gate 全部通过；crash/数据/权限/性能指标在目标阈值内。"),
        ],
        header_fill=BLUE_GRAY,
        first_col_bold=True,
        font_size=8.9,
    )
    add_heading(doc, "9.1 建议的工程阈值", 2)
    add_bullet(doc, "自动化回归：核心 package + app XCTest 通过；Release 版本也必须运行关键测试，不能只测 Debug。", bullet_id)
    add_bullet(doc, "稳定性：24 小时 soak 期间无未处理 fatal error、任务无界增长、内存趋势无持续泄漏；所有后台任务可取消。", bullet_id)
    add_bullet(doc, "交互：picker 从热键触发到首个可见 focus 的 p95 小于 150 ms；常见 FTS 查询 p95 小于 100 ms（以目标机器数据为准）。", bullet_id)
    add_bullet(doc, "数据：任何写入失败都有可见或可导出的状态；FTS 缺失可检测；backup 能在干净目录 restore 并通过校验。", bullet_id)
    add_bullet(doc, "隐私：自动化测试确认被过滤的 pasteboard 不进入 DB/blob/FTS；日志 fixture 不包含原文和秘密。", bullet_id)

    add_heading(doc, "10. Public release gate checklist", 1)
    add_text_paragraph(doc, "以下清单应作为发布 PR 的强制门禁。每一项都需要“通过证据”而不是口头确认；未通过项必须标记 owner、风险、临时缓解和预计关闭日期。")
    checklist = [
        ("[ ]", "构建", "Release 构建可重复，版本号、最低系统版本、架构和 entitlements 正确。"),
        ("[ ]", "签名", "Developer ID Application 签名；TeamIdentifier 正确；无 ad-hoc 签名残留。"),
        ("[ ]", "Hardened Runtime", "主 app、login item、extension 的 runtime/entitlements 经检查，最小权限原则成立。"),
        ("[ ]", "公证", "notarytool submit 成功；staple 成功；Gatekeeper 在干净机器上通过。"),
        ("[ ]", "权限", "首次安装、更新、权限撤回、重新授权、卸载/重装的 Accessibility/Input Monitoring/Keychain 行为可预期。"),
        ("[ ]", "数据", "quick_check、backup、restore、迁移回滚、FTS rebuild、blob 修复和 retention 强杀演练通过。"),
        ("[ ]", "输入", "单 representation、事件总量、文件/图片、未知 UTI、压缩炸弹式输入有硬上限和测试。"),
        ("[ ]", "粘贴", "普通 app、Terminal、浏览器、Electron、远程桌面、非 QWERTY layout、权限拒绝和目标 app 崩溃均有结果。"),
        ("[ ]", "交互", "⌘⇧V、focus、preview、arrow、Enter、Esc、multi-display、screen change、Reduce Motion 通过 smoke。"),
        ("[ ]", "隐私", "concealed/transient/auto-generated/排除 app 的默认行为、Capture All、日志脱敏和删除语义已文档化。"),
        ("[ ]", "观测", "崩溃报告/诊断包/结构化日志有 schema、保留策略、脱敏审核和用户可见的支持流程。"),
        ("[ ]", "回滚", "可下载上一稳定版本；数据库 schema 向后兼容或有迁移回滚策略；发布失败有撤回方案。"),
    ]
    add_table(doc, ["", "门禁", "验收"], checklist, [720, 1800, 6840], header_fill=LIGHT_GRAY, font_size=9.1, first_col_bold=False)
    add_heading(doc, "10.1 App Store / 直接分发注意事项", 2)
    add_bullet(doc, "如果选择 Developer ID + DMG/ZIP 直接分发，应把隐私说明、权限用途、签名身份和更新机制做完整；如果未来走 Mac App Store，需重新评估 Accessibility、login item、后台能力与沙盒限制。", bullet_id)
    add_bullet(doc, "不要把“本地存储”写成绝对隐私承诺；应明确本机存储、Keychain、诊断数据、备份导出和可选网络能力的边界。", bullet_id)
    add_bullet(doc, "更新器、签名、公证、卸载和登录项都属于产品体验的一部分，不应只当作发布脚本。", bullet_id)

    add_heading(doc, "11. 未来差异化方向", 1)
    add_text_paragraph(doc, "在 P0/P1 完成后，MacClippy 有空间形成区别于 Paste、Maccy 和 PasteBar 的产品表达。优先选择能复用已有原生优势、又不会牺牲隐私与稳定性的方向。")
    add_table(
        doc,
        ["方向", "用户价值", "实现约束"],
        [2300, 3560, 3500],
        rows=[
            ("可信的本地历史", "用户知道每条记录为何被捕获、保存多久、如何恢复和彻底删除。", "必须先有隐私默认、数据不变式、可观测性和恢复工具。"),
            ("面向工作流的 pinboard", "项目/客户/任务各自拥有可检索的临时工作集，支持批量顺序粘贴。", "保留原生 picker 速度；批量操作要可撤销、可预览。"),
            ("智能转换但可控", "对 URL、Markdown、JSON、代码、路径、富文本做本地转换/清理。", "所有 transform 明示输入/输出；不能暗中上传；失败保留原文。"),
            ("可解释的历史搜索", "结构化条件、来源 app、时间、类型、pinboard 和大小可组合筛选。", "FTS/metadata 索引必须可修复；搜索降级状态需可见。"),
            ("自动化接口", "CLI/MCP/Shortcuts 等接口服务高阶用户。", "脚本执行要有权限、审批、超时和数据范围限制；默认不执行任意脚本。"),
            ("可迁移与可备份", "用户不会被单一应用锁定，能安全带走历史。", "导出格式版本化、敏感数据明确、restore 需要完整校验。"),
        ],
        header_fill=BLUE_GRAY,
        first_col_bold=True,
        font_size=9,
    )
    add_callout(doc, "产品定位建议", "把 MacClippy 讲成“可信的原生 Mac 工作记忆”，而不是“又一个剪贴板历史”。关键词应是 fast、local、recoverable、keyboard-first；所有智能能力都应服务于找回和复用，而不是制造不可预测的自动化。", fill=BLUE_GRAY)

    add_heading(doc, "12. 附录：关键代码位置与参考资料", 1)
    add_heading(doc, "12.1 建议优先阅读/修改的位置", 2)
    code_items = [
        ("运行时与生命周期", "/Users/mingjie.wang/Documents/personal/mac-all-you-need/mac-clippy/MacClippy/MacClippyRuntime.swift:263"),
        ("pasteboard 捕获", "/Users/mingjie.wang/Documents/personal/mac-all-you-need/mac-clippy/MacClippyKit/Sources/MacClippyPlatform/PasteboardCapture.swift"),
        ("reconciliation / startup repair", "/Users/mingjie.wang/Documents/personal/mac-all-you-need/mac-clippy/MacClippyKit/Sources/MacClippyCore/MacClippyReconciliation.swift"),
        ("retention 与清理", "/Users/mingjie.wang/Documents/personal/mac-all-you-need/mac-clippy/MacClippyKit/Sources/MacClippyCore/RetentionPolicy.swift"),
        ("launch at login", "/Users/mingjie.wang/Documents/personal/mac-all-you-need/mac-clippy/MacClippy/MacClippyApp.swift:84"),
        ("构建、签名与测试入口", "/Users/mingjie.wang/Documents/personal/mac-all-you-need/mac-clippy/project.yml；/Users/mingjie.wang/Documents/personal/mac-all-you-need/mac-clippy/Makefile"),
    ]
    for label, path in code_items:
        p = doc.add_paragraph(style="Code Paragraph")
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(label + "  ")
        set_run_font(r, FONT, 9, INK, True, False)
        r = p.add_run(path)
        set_run_font(r, MONO_FONT, 8.5, DARK_BLUE, False, False)

    add_heading(doc, "12.2 参考链接", 2)
    refs = [
        ("Paste 官方产品页", "https://pasteapp.io/"),
        ("Deck DeepWiki", "https://deepwiki.com/yuzeguitarist/Deck"),
        ("Maccy DeepWiki", "https://deepwiki.com/p0deje/Maccy"),
        ("Clipy DeepWiki", "https://deepwiki.com/Clipy/Clipy"),
        ("PasteBar DeepWiki", "https://deepwiki.com/PasteBar/PasteBarApp"),
        ("CopyQ GitHub", "https://github.com/hluk/CopyQ"),
        ("Tauri architecture", "https://v2.tauri.app/concept/architecture/"),
        ("Tauri macOS signing", "https://v2.tauri.app/distribute/sign/macos/"),
        ("Maccy GitHub", "https://github.com/p0deje/Maccy"),
        ("PasteBarApp GitHub", "https://github.com/PasteBar/PasteBarApp"),
    ]
    for label, url in refs:
        p = doc.add_paragraph(style="Body Small")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(label + "：")
        set_run_font(r, FONT, 9.2, "222222", True, False)
        add_hyperlink(p, url, url)

    add_heading(doc, "12.3 最终建议", 2)
    add_text_paragraph(doc, "先把 MacClippy 从“功能已完成”推进到“发布和恢复已证明”：先签名、公证、数据恢复、输入边界、并发和隐私默认，再做更大范围的跨平台或智能功能扩张。完成 P0/P1 后，MacClippy 才真正具备让用户长期托付剪贴板历史的基础。")
    add_callout(doc, "Go / No-Go", "当前建议：No-Go for public production release；Go for focused hardening sprint and private/advanced-beta validation。P0 关闭并通过 release gate 后，再进入 public release candidate。", fill=PALE_GOLD, label_color=GOLD)

    # Core properties for document metadata.
    props = doc.core_properties
    props.title = "MacClippy 生产就绪性审计与产品化路线图"
    props.subject = "Swift vs. Tauri、竞品研究、稳定性改进与发布门禁"
    props.author = "Codex"
    props.keywords = "MacClippy, macOS, clipboard, production readiness, Swift, Tauri"
    props.comments = "Prepared for mingjie-father"

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
